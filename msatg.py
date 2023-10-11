from docx import Document, shared
from datetime import date
import calendar
import os, sys, subprocess

def open_docx(fname):
    if sys.platform == 'win32':
        os.startfile(fname)
    else:
        opener = 'open' if sys.platform == 'darwin' else 'xdg-open'
        subprocess.call([opener, fname])

# preferences
NAME = 'Nic Bolton'
STUDENT_ID = '110034690'
FORCE_CAPITALIZE = True
FONT = 'Times New Roman'
FONT_SIZE = 12

# input
doc_title = input('File name: ')

course = input('Course code: ')
course = course.upper() if FORCE_CAPITALIZE else course

asgn_title = input('Title: ')
asgn_title = asgn_title.title() if FORCE_CAPITALIZE else asgn_title
    
today = date.today().strftime("%d %m %Y")
month = calendar.month_name[int(today[3:5])]
today = today[:2] + ' ' + month + ' ' + today[6:]

# create file if it doesn't exist
doc = Document()

# set styles
style = doc.styles['Normal']
font = style.font
font.name = FONT
font.size = shared.Pt(FONT_SIZE)

p = doc.add_paragraph()
p.add_run(course + ' ' + asgn_title + '\n').bold = True
p.add_run(NAME + ' #' + STUDENT_ID + '\n').bold = True
p.add_run(today).bold = True

doc.add_paragraph('\n')

doc.save(doc_title + '.docx')

open_docx(doc_title + '.docx')

